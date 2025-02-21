from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
import requests
import shutil
from pathlib import Path
from docx import Document
import os
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

# Initialize FastAPI application
app = FastAPI()

# Enable CORS for frontend connection
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allow all origins
    allow_credentials=True,
    allow_methods=["*"],  # Allow all methods
    allow_headers=["*"],  # Allow all headers
)

# Directory to store uploaded & processed files
UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)

# Retrieve OpenAI API Key from .env file
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

# OpenAI API endpoint for chat-based text completion
OPENAI_URL = "https://api.openai.com/v1/chat/completions"  # Correct endpoint

# Ensure the API Key exists before making requests
if not OPENAI_API_KEY:
    raise ValueError("Missing OpenAI API Key. Please check your .env file.")

# Get character limit for a full-width line
def get_full_line_char_limit(doc):
    return max(len(p.text) for p in doc.paragraphs if len(p.text) > 50)

# Call OpenAI API to expand short bullet points
def expand_text_with_openai(text, additional_chars_needed):
    messages = [
        {"role": "system", "content": "You are an assistant that expands short sentences while keeping them concise."},
        {"role": "user", "content": f"Expand this sentence by approximately {additional_chars_needed} characters: {text}"}
    ]
    
    print(f"Sending request to OpenAI: {messages[1]['content']}")  # Debugging output

    response = requests.post(
        OPENAI_URL,
        headers={"Authorization": f"Bearer {OPENAI_API_KEY}"},
        json={
            "model": "gpt-3.5-turbo",
            "messages": messages,
            "max_tokens": additional_chars_needed,
            "temperature": 0.7
        }
    )

    if response.status_code == 200:
        expanded_text = response.json().get("choices", [{}])[0].get("message", {}).get("content", text)
        print(f"Received expanded text: {expanded_text}")  # Debugging output
        return expanded_text.strip()
    else:
        print(f"OpenAI API Error: {response.text}")  # Print error details
        raise HTTPException(status_code=response.status_code, detail="Error with OpenAI API request")

# Process and expand the smallest bullet point in the .docx resume
def process_and_expand_smallest_bullet_point(file_path):
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"Error opening document: {str(e)}")  # Log the error
        raise HTTPException(status_code=500, detail=f"Error opening document: {str(e)}")

    full_line_char_limit = get_full_line_char_limit(doc)
    print(f"Full line character limit: {full_line_char_limit}")  # Debugging output

    smallest_bullet_point = None
    smallest_length = float('inf')  # Initially set to infinity

    # Find the smallest bullet point
    for para in doc.paragraphs:
        print(f"Checking paragraph: {para.text}")  # Debugging output
        if para.style and para.style.name.startswith('List'):
            print(f"Bullet point detected: {para.text}")  # Debugging output
            line_length = len(para.text)
            if line_length < smallest_length:
                smallest_bullet_point = para
                smallest_length = line_length
                print(f"Smallest bullet point found: {para.text}")  # Debugging output

    # Expand the smallest bullet point if found
    if smallest_bullet_point:
        line_length_diff = full_line_char_limit - smallest_length
        if line_length_diff > 15:  # Only expand if the difference is significant enough
            print(f"Expanding bullet point by {line_length_diff} characters")  # Debugging output
            try:
                expanded_text = expand_text_with_openai(smallest_bullet_point.text, line_length_diff)
                smallest_bullet_point.text = expanded_text  # Replace with expanded version
            except Exception as e:
                print(f"Error expanding bullet point: {str(e)}")  # Log expansion error
                raise HTTPException(status_code=500, detail=f"Error expanding text: {str(e)}")
        else:
            print("Bullet point is already long enough, no expansion needed.")  # Debugging output
    else:
        print("No bullet points found in the document.")  # Debugging output

    # Save the modified document
    new_file_path = file_path.parent / f"expanded_{file_path.name}"
    try:
        doc.save(new_file_path)  # Save the document with changes
    except Exception as e:
        print(f"Error saving the file: {str(e)}")  # Log saving error
        raise HTTPException(status_code=500, detail=f"Error saving the file: {str(e)}")

    print(f"Processed file saved at {new_file_path}")  # Debugging output

    # Display modified document content
    print("\nModified Document Content:")
    for para in doc.paragraphs:
        print(f"Paragraph: {para.text}")

    return new_file_path

# Upload endpoint - accepts .docx file, expands it, and returns a download link
@app.post("/upload/")
async def upload_and_process_file(file: UploadFile = File(...)):
    try:
        file_location = UPLOAD_DIR / file.filename
        with file_location.open("wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        print(f"File uploaded: {file.filename}")  # Debugging output

        expanded_file_path = process_and_expand_smallest_bullet_point(file_location)
        return {"download_url": f"http://127.0.0.1:3001/download/{expanded_file_path.name}"}
    except Exception as e:
        print(f"Error processing file: {str(e)}")  # Log the error
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")

# Download endpoint - allows users to download the expanded resume
@app.get("/download/{filename}")
async def download_file(filename: str):
    file_path = UPLOAD_DIR / filename
    if file_path.exists():
        return FileResponse(file_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=filename)
    raise HTTPException(status_code=404, detail="File not found")